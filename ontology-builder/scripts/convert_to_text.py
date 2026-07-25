#!/usr/bin/env python3
"""Convert an office document to plain text/Markdown so Claude can read it.

Claude Code's Read tool already handles .md, .txt and .pdf natively — this
script only needs to cover formats it can't read directly (.docx, .doc,
.odt, .rtf, .pptx, ...). Two conversion paths are tried, in order:

1. python-docx  — fast, no external dependency beyond the pip package,
   but only understands the modern .docx (Office Open XML) format.
2. pandoc       — external binary, covers many more formats (.doc, .odt,
   .rtf, .pptx...) if it happens to be installed on the machine.

Usage:
    python convert_to_text.py <input_file> [-o <output_file>]

If -o is omitted, the extracted text is printed to stdout so it can be
piped or captured directly without leaving a file behind.

Note on accented characters: the file this script writes is always valid
UTF-8. If you view its stdout output through a Windows terminal (Git Bash,
cmd.exe) and accented characters look garbled, that's the terminal's own
codepage misreading UTF-8 — not a bug in the conversion. Prefer -o to write
a file and read it back with a UTF-8-aware tool rather than trusting the
console rendering.
"""
from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from pathlib import Path


def convert_docx(path: Path) -> str:
    """Extract paragraphs and tables from a .docx, in document order where possible."""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError(
            "python-docx n'est pas installé. Installe-le avec : pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    parts: list[str] = []

    for element in document.element.body.iterchildren():
        tag = element.tag.rsplit("}", 1)[-1]
        if tag == "p":
            para = next((p for p in document.paragraphs if p._p is element), None)
            text = para.text.strip() if para is not None else ""
            if text:
                # Heading styles become Markdown headings so structure survives conversion.
                style = (para.style.name if para is not None and para.style else "") or ""
                if style.startswith("Heading"):
                    try:
                        level = int(style.split(" ")[-1])
                    except ValueError:
                        level = 2
                    parts.append(f"{'#' * min(level, 6)} {text}")
                else:
                    parts.append(text)
        elif tag == "tbl":
            table = next((t for t in document.tables if t._tbl is element), None)
            if table is not None:
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if rows:
                    parts.append(_rows_to_markdown_table(rows))

    return "\n\n".join(parts)


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    header, *body = rows
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * len(header)) + " |"]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def convert_with_pandoc(path: Path) -> str:
    if not shutil.which("pandoc"):
        raise RuntimeError(
            f"pandoc n'est pas disponible pour convertir {path.suffix} — "
            "installe pandoc (https://pandoc.org/installing.html) ou fournis "
            "une version .docx/.md/.pdf/.txt du document."
        )
    result = subprocess.run(
        ["pandoc", str(path), "-t", "markdown", "--wrap=none"],
        capture_output=True, text=True, check=True,
    )
    return result.stdout


def convert(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            return convert_docx(path)
        except RuntimeError:
            return convert_with_pandoc(path)
    return convert_with_pandoc(path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_file", type=Path)
    parser.add_argument("-o", "--output", type=Path, default=None)
    args = parser.parse_args()

    if not args.input_file.is_file():
        print(f"Fichier introuvable : {args.input_file}", file=sys.stderr)
        sys.exit(1)

    try:
        text = convert(args.input_file)
    except RuntimeError as exc:
        print(str(exc), file=sys.stderr)
        sys.exit(1)

    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text, encoding="utf-8")
        print(f"Écrit : {args.output}")
    else:
        print(text)


if __name__ == "__main__":
    main()
